from __future__ import annotations

import hashlib
import json
import mimetypes
import re
import shutil
import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
from docx import Document as DocxDocument
from pypdf import PdfReader

from .config import EnterpriseConfig
from .database import Database, utcnow
from .providers import EmbeddingProvider
from .security import Principal, detect_prompt_injection, safe_component, safe_join, scope_clause
from .vector_store import VectorStore


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()


@dataclass
class ExtractedChunk:
    text: str
    page: Optional[int] = None
    sheet: Optional[str] = None
    section: Optional[str] = None
    row_range: Optional[str] = None
    location_key: str = ""
    metadata: Optional[Dict[str, Any]] = None


class DocumentService:
    def __init__(self, cfg: EnterpriseConfig, db: Database, embeddings: EmbeddingProvider, vectors: VectorStore, dataset_service=None):
        self.cfg = cfg
        self.db = db
        self.embeddings = embeddings
        self.vectors = vectors
        self.dataset_service = dataset_service

    def _validate(self, path: Path) -> None:
        doc_cfg = self.cfg.section("documents")
        if path.suffix.lower() not in set(doc_cfg.get("allowed_extensions", [])):
            raise ValueError(f"Formato no permitido: {path.suffix.lower()}")
        max_bytes = int(doc_cfg.get("max_file_mb", 250)) * 1024 * 1024
        if path.stat().st_size > max_bytes:
            raise ValueError(f"Archivo excede {doc_cfg.get('max_file_mb')} MB")

    def _split_text(self, text: str, prefix: ExtractedChunk) -> List[ExtractedChunk]:
        size = int(self.cfg.section("documents").get("chunk_chars", 1200))
        overlap = int(self.cfg.section("documents").get("chunk_overlap", 180))
        clean = re.sub(r"\x00", "", text or "").strip()
        if not clean:
            return []
        output: List[ExtractedChunk] = []
        start = 0
        index = 0
        while start < len(clean):
            end = min(len(clean), start + size)
            if end < len(clean):
                cut = max(clean.rfind("\n", start, end), clean.rfind(". ", start, end))
                if cut > start + size // 2:
                    end = cut + 1
            part = clean[start:end].strip()
            if part:
                output.append(
                    ExtractedChunk(
                        part, prefix.page, prefix.sheet, prefix.section, prefix.row_range,
                        f"{prefix.location_key}:part{index}", dict(prefix.metadata or {}),
                    )
                )
            if end >= len(clean):
                break
            start = max(start + 1, end - overlap)
            index += 1
        return output

    def _extract_tabular(self, path: Path) -> Tuple[List[ExtractedChunk], Dict[str, Any]]:
        max_rows = int(self.cfg.section("documents").get("tabular_preview_rows", 300))
        chunks: List[ExtractedChunk] = []
        info: List[Dict[str, Any]] = []
        if path.suffix.lower() == ".csv":
            frames = {"CSV": pd.read_csv(path, nrows=max_rows, low_memory=False)}
        else:
            try:
                xls = pd.ExcelFile(path, engine="calamine")
            except Exception:
                xls = pd.ExcelFile(path)
            frames = {}
            try:
                for sheet in xls.sheet_names:
                    try:
                        frames[str(sheet)] = pd.read_excel(xls, sheet_name=sheet, nrows=max_rows)
                    except Exception as exc:
                        info.append({"sheet": str(sheet), "error": str(exc)})
            finally:
                try:
                    xls.close()
                except Exception:
                    pass
        for sheet, frame in frames.items():
            columns = [str(c) for c in frame.columns]
            info.append({"sheet": sheet, "columns": columns, "preview_rows": int(len(frame))})
            schema = f"Dataset tabular. Hoja: {sheet}. Columnas: {', '.join(columns)}. Tipos aproximados: " + ", ".join(f"{c}={frame[c].dtype}" for c in frame.columns)
            chunks.append(ExtractedChunk(schema, sheet=sheet, section="Esquema", location_key=f"sheet:{sheet}:schema", metadata={"tabular_schema": True}))
            for start in range(0, min(len(frame), max_rows), 50):
                subset = frame.iloc[start:start + 50]
                lines = []
                for _, row in subset.iterrows():
                    values = []
                    for column in frame.columns[:20]:
                        value = row[column]
                        if pd.isna(value):
                            continue
                        values.append(f"{column}={str(value)[:120]}")
                    if values:
                        lines.append("; ".join(values))
                if lines:
                    row_range = f"{start + 2}-{start + 1 + len(subset)}"
                    chunks += self._split_text(
                        "\n".join(lines),
                        ExtractedChunk("", sheet=sheet, section="Muestra tabular", row_range=row_range, location_key=f"sheet:{sheet}:rows:{row_range}", metadata={"tabular_preview": True}),
                    )
        return chunks, {"type": "tabular", "sheets": info, "numeric_calculations": "structured_tool_only"}

    def _extract(self, path: Path) -> Tuple[List[ExtractedChunk], Dict[str, Any]]:
        ext = path.suffix.lower()
        if ext == ".pdf":
            reader = PdfReader(str(path))
            chunks: List[ExtractedChunk] = []
            for page_no, page in enumerate(reader.pages, 1):
                chunks += self._split_text(page.extract_text() or "", ExtractedChunk("", page=page_no, location_key=f"page:{page_no}"))
            return chunks, {"type": "pdf", "pages": len(reader.pages)}
        if ext == ".docx":
            doc = DocxDocument(str(path))
            chunks = []
            section_no = 1
            current: List[str] = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                if paragraph.style and str(paragraph.style.name).lower().startswith("heading") and current:
                    chunks += self._split_text("\n".join(current), ExtractedChunk("", section=f"Seccion {section_no}", location_key=f"section:{section_no}"))
                    current = []
                    section_no += 1
                current.append(text)
            if current:
                chunks += self._split_text("\n".join(current), ExtractedChunk("", section=f"Seccion {section_no}", location_key=f"section:{section_no}"))
            for table_no, table in enumerate(doc.tables, 1):
                rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
                chunks += self._split_text("\n".join(rows), ExtractedChunk("", section=f"Tabla {table_no}", location_key=f"table:{table_no}"))
            return chunks, {"type": "docx", "sections": section_no, "tables": len(doc.tables)}
        if ext in {".txt", ".md", ".markdown"}:
            return self._split_text(path.read_text(encoding="utf-8", errors="replace"), ExtractedChunk("", section="Documento", location_key="text")), {"type": "text"}
        if ext in {".csv", ".xlsx", ".xls", ".xlsm", ".xlsb"}:
            return self._extract_tabular(path)
        raise ValueError("Formato no soportado")

    def index(self, principal: Principal, source_path: str | Path, *, scope: str = "company", display_name: Optional[str] = None, force: bool = False) -> Dict[str, Any]:
        source = Path(source_path).resolve()
        self._validate(source)
        name = safe_component(display_name or source.name)
        file_hash = sha256_file(source)
        owner = principal.user_id if scope == "user" else None
        existing = self.db.one(
            "SELECT * FROM documents WHERE company_id=? AND scope=? AND ((user_id IS NULL AND ? IS NULL) OR user_id=?) AND name=?",
            (principal.company_id, scope, owner, owner, name),
        )
        if existing and existing["current_hash"] == file_hash and existing["active"] and not force:
            return {"ok": True, "unchanged": True, "document_id": existing["id"], "version": existing["current_version"], "name": name}
        document_id = existing["id"] if existing else str(uuid.uuid4())
        version = int(existing["current_version"]) + 1 if existing else 1
        destination_dir = safe_join(self.cfg.knowledge_dir, principal.company_id, document_id)
        destination_dir.mkdir(parents=True, exist_ok=True)
        destination = destination_dir / f"v{version}_{name}"
        shutil.copy2(source, destination)
        chunks, metadata = self._extract(destination)
        now = utcnow()
        new_vector_ids: List[str] = []
        missing_chunks: List[Tuple[str, ExtractedChunk]] = []
        for chunk in chunks:
            text_hash = sha256_text(chunk.text)
            vector_id = str(uuid.uuid5(uuid.NAMESPACE_URL, f"{document_id}|{chunk.location_key}|{text_hash}"))
            new_vector_ids.append(vector_id)
            if not self.vectors.has("document", vector_id):
                missing_chunks.append((vector_id, chunk))
        if missing_chunks:
            vectors = self.embeddings.embed([chunk.text for _, chunk in missing_chunks])
            for (vector_id, chunk), vector in zip(missing_chunks, vectors):
                self.vectors.upsert(
                    "document", vector_id, vector,
                    {
                        "document_id": document_id, "company_id": principal.company_id, "user_id": owner,
                        "scope": scope, "name": name, "page": chunk.page, "sheet": chunk.sheet,
                        "section": chunk.section, "row_range": chunk.row_range,
                        "injection_flag": detect_prompt_injection(chunk.text),
                    },
                )
        old_vector_ids = [row["vector_id"] for row in self.db.query("SELECT vector_id FROM document_chunks WHERE document_id=? AND active=1", (document_id,))]
        with self.db.tx() as con:
            con.execute("UPDATE document_chunks SET active=0 WHERE document_id=?", (document_id,))
            if existing:
                con.execute(
                    "UPDATE documents SET stored_path=?,extension=?,mime_type=?,size_bytes=?,current_hash=?,current_version=?,status='ready',active=1,updated_at=? WHERE id=?",
                    (str(destination), source.suffix.lower(), mimetypes.guess_type(name)[0], source.stat().st_size, file_hash, version, now, document_id),
                )
            else:
                con.execute(
                    "INSERT INTO documents(id,company_id,user_id,scope,name,stored_path,extension,mime_type,size_bytes,current_hash,current_version,status,active,added_by,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                    (document_id, principal.company_id, owner, scope, name, str(destination), source.suffix.lower(), mimetypes.guess_type(name)[0], source.stat().st_size, file_hash, version, "ready", 1, principal.user_id, now, now),
                )
            con.execute(
                "INSERT INTO document_versions(id,document_id,version,file_hash,stored_path,indexed_at,chunk_count,metadata_json) VALUES(?,?,?,?,?,?,?,?)",
                (str(uuid.uuid4()), document_id, version, file_hash, str(destination), now, len(chunks), json.dumps(metadata, ensure_ascii=False)),
            )
            for ordinal, chunk in enumerate(chunks):
                text_hash = sha256_text(chunk.text)
                vector_id = str(uuid.uuid5(uuid.NAMESPACE_URL, f"{document_id}|{chunk.location_key}|{text_hash}"))
                con.execute(
                    "INSERT INTO document_chunks(id,vector_id,document_id,version,ordinal,text_hash,content,page,sheet,section,row_range,injection_flag,active,metadata_json,created_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                    (
                        str(uuid.uuid4()), vector_id, document_id, version, ordinal, text_hash, chunk.text,
                        chunk.page, chunk.sheet, chunk.section, chunk.row_range,
                        1 if detect_prompt_injection(chunk.text) else 0, 1,
                        json.dumps(chunk.metadata or {}, ensure_ascii=False), now,
                    ),
                )
        removed = set(old_vector_ids) - set(new_vector_ids)
        self.vectors.delete("document", removed)
        if self.dataset_service and source.suffix.lower() in {".csv", ".xlsx", ".xls", ".xlsm", ".xlsb"}:
            try:
                self.dataset_service.register(principal, destination, name=name, scope=scope, file_hash=file_hash)
            except Exception as exc:
                metadata["dataset_registration_warning"] = str(exc)
        self.db.audit(
            "document.index", principal.company_id, principal.user_id, "document", document_id,
            details={"name": name, "version": version, "chunks": len(chunks), "new_embeddings": len(missing_chunks), "removed_vectors": len(removed)},
        )
        return {"ok": True, "unchanged": False, "document_id": document_id, "version": version, "name": name, "hash": file_hash, "chunks": len(chunks), "new_embeddings": len(missing_chunks), "metadata": metadata}

    def list(self, principal: Principal, include_inactive: bool = False) -> List[Dict[str, Any]]:
        clause, args = scope_clause(principal)
        sql = f"SELECT * FROM documents WHERE {clause}" + ("" if include_inactive else " AND active=1") + " ORDER BY updated_at DESC"
        return [dict(row) for row in self.db.query(sql, args)]

    def search(self, principal: Principal, query: str, limit: int = 8, min_score: float = 0.18) -> List[Dict[str, Any]]:
        if not query.strip():
            return []
        # Evita generar embeddings si todavía no hay documentos accesibles.
        clause, args = scope_clause(principal, "d")
        exists = self.db.one(f"SELECT 1 FROM documents d WHERE d.active=1 AND {clause} LIMIT 1", args)
        if not exists:
            return []
        query_vector = self.embeddings.embed([query])[0]
        hits = self.vectors.search("document", query_vector, principal, max(limit * 4, 20))
        output = []
        for hit in hits:
            row = self.db.one(
                "SELECT c.*,d.name,d.current_version,d.company_id,d.user_id,d.scope,d.active AS doc_active FROM document_chunks c JOIN documents d ON d.id=c.document_id WHERE c.vector_id=? AND c.active=1 AND d.active=1 AND d.current_version=c.version",
                (hit["vector_id"],),
            )
            if not row:
                continue
            if row["company_id"] != principal.company_id or (row["scope"] != "company" and row["user_id"] != principal.user_id):
                continue
            if float(hit["score"]) < min_score:
                continue
            data = dict(row)
            data["score"] = round(float(hit["score"]), 4)
            output.append(data)
            if len(output) >= limit:
                break
        return output

    def delete(self, principal: Principal, document_id: str) -> None:
        clause, args = scope_clause(principal, "d")
        row = self.db.one(f"SELECT d.* FROM documents d WHERE d.id=? AND {clause}", (document_id, *args))
        if not row:
            raise KeyError("Documento no encontrado")
        vector_ids = [r["vector_id"] for r in self.db.query("SELECT vector_id FROM document_chunks WHERE document_id=? AND active=1", (document_id,))]
        self.vectors.delete("document", vector_ids)
        self.db.execute("UPDATE documents SET active=0,status='deleted',updated_at=? WHERE id=?", (utcnow(), document_id))
        self.db.execute("UPDATE document_chunks SET active=0 WHERE document_id=?", (document_id,))
        self.db.audit("document.delete", principal.company_id, principal.user_id, "document", document_id, details={"name": row["name"]})

    def reindex(self, principal: Principal, document_id: str) -> Dict[str, Any]:
        clause, args = scope_clause(principal, "d")
        row = self.db.one(f"SELECT d.* FROM documents d WHERE d.id=? AND {clause}", (document_id, *args))
        if not row:
            raise KeyError("Documento no encontrado")
        return self.index(principal, row["stored_path"], scope=row["scope"], display_name=row["name"], force=True)
